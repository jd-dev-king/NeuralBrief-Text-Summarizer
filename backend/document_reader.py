"""Read text from TXT, DOCX, and PDF documents."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader


class DocumentReader:
    """Extract readable text from supported document formats."""

    SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}

    def read(self, file_path: str | Path) -> str:
        """
        Read text from a supported document.

        Args:
            file_path: Path to the uploaded document.

        Returns:
            Extracted document text.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file type is unsupported or contains no text.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                "Unsupported file type. Please upload a TXT, DOCX, or PDF file."
            )

        if extension == ".txt":
            text = self._read_txt(path)
        elif extension == ".docx":
            text = self._read_docx(path)
        else:
            text = self._read_pdf(path)

        cleaned_text = self._clean_text(text)

        if not cleaned_text:
            raise ValueError(
                "No readable text was found in the uploaded document."
            )

        return cleaned_text

    @staticmethod
    def _read_txt(path: Path) -> str:
        """Read a plain-text file with common encoding fallbacks."""
        encodings = ("utf-8", "utf-8-sig", "latin-1")

        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue

        raise ValueError(
            "The text file could not be decoded using a supported encoding."
        )

    @staticmethod
    def _read_docx(path: Path) -> str:
        """Extract paragraphs and table text from a Word document."""
        try:
            document = Document(path)
        except Exception as error:
            raise ValueError(
                "The Word document could not be opened."
            ) from error

        text_sections: list[str] = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                text_sections.append(paragraph_text)

        for table in document.tables:
            for row in table.rows:
                row_values = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if row_values:
                    text_sections.append(" | ".join(row_values))

        return "\n".join(text_sections)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        """Extract text from every readable page in a PDF."""
        try:
            reader = PdfReader(path)
        except Exception as error:
            raise ValueError(
                "The PDF document could not be opened."
            ) from error

        page_text: list[str] = []

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                extracted_text = page.extract_text()
            except Exception:
                extracted_text = None

            if extracted_text and extracted_text.strip():
                page_text.append(extracted_text.strip())

        if not page_text:
            raise ValueError(
                "No readable text was found in the PDF. "
                "Scanned image-only PDFs are not currently supported."
            )

        return "\n".join(page_text)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Remove null characters and excessive whitespace."""
        lines = [
            " ".join(line.replace("\x00", " ").split())
            for line in text.splitlines()
        ]

        return "\n".join(
            line for line in lines if line
        ).strip()